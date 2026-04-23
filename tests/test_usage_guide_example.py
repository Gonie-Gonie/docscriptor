from __future__ import annotations

import importlib.util
from io import BytesIO
from pathlib import Path

from docx import Document as WordDocument
from pypdf import PdfReader


def _load_example_module(example_dir: str):
    module_path = Path(__file__).resolve().parents[1] / "examples" / example_dir / "main.py"
    spec = importlib.util.spec_from_file_location(f"examples.{example_dir}.main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _pdf_image_draw_count(pdf_path: Path) -> int:
    count = 0
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        resources = page.get("/Resources")
        if resources is None or "/XObject" not in resources:
            continue
        xobjects = resources["/XObject"].get_object()
        image_names = {
            name
            for name, xobject in xobjects.items()
            if xobject.get_object().get("/Subtype") == "/Image"
        }
        if not image_names:
            continue
        content = page.get_contents()
        if content is None:
            continue
        content_bytes = content.get_data()
        for name in image_names:
            token = f"{name} Do".encode()
            count += content_bytes.count(token)
    return count


def test_usage_guide_example_builds_outputs(tmp_path: Path) -> None:
    usage_guide = _load_example_module("usage_guide_example")
    docx_path, pdf_path = usage_guide.build_usage_guide(tmp_path)

    assert docx_path.exists()
    assert pdf_path.exists()
    assert (Path(usage_guide.__file__).resolve().parent / "assets" / "usage-guide-figure.png").exists()

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)

    assert "Using docscriptor" in paragraph_texts
    assert "List of Tables" in paragraph_texts
    assert "List of Figures" in paragraph_texts
    assert "Contents" in paragraph_texts
    assert "Footnotes" in paragraph_texts
    assert "Comments" in paragraph_texts
    assert "References" in paragraph_texts
    assert "1 Getting Started" in paragraph_texts
    assert "1.1 Quick Start" in paragraph_texts
    assert "2 Authoring Model" in paragraph_texts
    assert "2.1 Core Blocks" in paragraph_texts
    assert "2.2 Inline Actions" in paragraph_texts
    assert "2.3 Reusable Patterns" in paragraph_texts
    assert any("Text.bold(...)" in text for text in paragraph_texts)
    assert any("Text.from_markup(...)" in text for text in paragraph_texts)
    assert any("The project repository can be cited directly as [1]" in text for text in paragraph_texts)
    assert any("generated comments page" in text for text in paragraph_texts)
    assert any("Portable footnotes remain stable" in text for text in paragraph_texts)
    assert any("dx = (" in text and ")/(3)" in text for text in paragraph_texts)
    assert any("Keep reusable assets in an asset directory beside the example script." in text for text in paragraph_texts)
    assert word_document.sections[0].footer.paragraphs[0].text.startswith("Page ")
    assert len(word_document.tables) == 2
    assert len(word_document.inline_shapes) == 1
    assert paragraph_texts.count("Table 1. Rendering outputs by goal.") >= 2
    assert paragraph_texts.count("Table 2. Core authoring primitives.") >= 2
    assert paragraph_texts.count("Figure 1. Example figure loaded directly from the example asset directory.") >= 2

    assert "Using docscriptor" in pdf_text
    assert "List of Tables" in pdf_text
    assert "List of Figures" in pdf_text
    assert "Contents" in pdf_text
    assert "Footnotes" in pdf_text
    assert "Comments" in pdf_text
    assert "References" in pdf_text
    assert "Getting Started" in pdf_text
    assert "Authoring Model" in pdf_text
    assert "Quick Start" in pdf_text
    assert "Core Blocks" in pdf_text
    assert "Inline Actions" in pdf_text
    assert "Reusable Patterns" in pdf_text
    assert "The project repository can be cited directly as [1]" in pdf_text
    assert "Portable footnotes remain stable" in pdf_text
    assert "Text.bold(...)" in pdf_text
    assert "Text.from_markup(...)" in pdf_text
    assert "Literate Programming" in pdf_text
    assert "https://github.com/Gonie-Gonie/pydocs" in pdf_text
    assert _pdf_image_draw_count(pdf_path) == 1
